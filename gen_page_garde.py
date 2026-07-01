from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import copy

doc = Document()

# ── Marges ────────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width    = Cm(21)
sec.page_height   = Cm(29.7)
sec.top_margin    = Cm(1.8)
sec.bottom_margin = Cm(1.5)
sec.left_margin   = Cm(2)
sec.right_margin  = Cm(2)

# ── Bordure double autour de la page ─────────────────────────────────────────
def set_page_border(doc):
    sectPr = doc.sections[0]._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    pgBorders.set(qn('w:display'), 'allPages')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'double')
        b.set(qn('w:sz'), '18')
        b.set(qn('w:space'), '24')
        b.set(qn('w:color'), '1F3A8A')
        pgBorders.append(b)
    sectPr.append(pgBorders)

set_page_border(doc)

# ── Helpers ───────────────────────────────────────────────────────────────────
BLEU   = RGBColor(0x1F, 0x3A, 0x8A)
ORANGE = RGBColor(0xC0, 0x50, 0x00)
BLANC  = RGBColor(0xFF, 0xFF, 0xFF)
NOIR   = RGBColor(0x00, 0x00, 0x00)

def remove_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def para_spacing(p, before=0, after=40):
    pf = p.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after  = Twips(after)
    pf.line_spacing = Pt(13)

def add_run(p, text, size=10, bold=False, italic=False, color=None, all_caps=False):
    r = p.add_run(text)
    r.font.size  = Pt(size)
    r.bold       = bold
    r.italic     = italic
    r.font.all_caps = all_caps
    if color:
        r.font.color.rgb = color
    return r

def simple_para(doc, text, size=10, bold=False, italic=False,
                align=WD_ALIGN_PARAGRAPH.CENTER, color=None,
                before=0, after=40, all_caps=False):
    p = doc.add_paragraph()
    p.alignment = align
    para_spacing(p, before, after)
    add_run(p, text, size, bold, italic, color, all_caps)
    return p

def separator_para(doc, before=20, after=20):
    p = doc.add_paragraph('— — — — — — — — — — — — — — — —')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before, after)
    r = p.runs[0]
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x44,0x44,0x44)
    return p

def tiret_para(doc, before=10, after=10):
    p = doc.add_paragraph('————————')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, before, after)
    r = p.runs[0]
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0x44,0x44,0x44)
    return p

# ── 1. EN-TÊTE BILINGUE (tableau 2 colonnes) ─────────────────────────────────
tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

left_lines = [
    ("REPUBLIQUE DU CAMEROUN", 8.5, False),
    ("Paix-Travail-Patrie",    8,   True),
    ("————————",               7,   False),
    ("MINISTERE DE",           8.5, False),
    ("L'ENSEIGNEMENT SUPERIEUR", 8.5, False),
    ("————————",               7,   False),
    ("INSTITUT SUPERIEUR DE",  8.5, False),
    ("TECHNOLOGIE APPLIQUEE",  8.5, False),
    ("ET DE GESTION",          8.5, False),
]
right_lines = [
    ("REPUBLIC OF CAMEROON",   8.5, False),
    ("Peace-Work-Fatherland",  8,   True),
    ("————————",               7,   False),
    ("MINISTRY OF HIGHER",     8.5, False),
    ("EDUCATION",              8.5, False),
    ("————————",               7,   False),
    ("HIGHER INSTITUTE OF APPLIED", 8.5, False),
    ("TECHNOLOGY AND",         8.5, False),
    ("MANAGEMENT",             8.5, False),
]

for col_idx, lines in enumerate([left_lines, right_lines]):
    cell = tbl.rows[0].cells[col_idx]
    remove_cell_borders(cell)
    for i, (txt, sz, it) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, 0, 16)
        r = p.add_run(txt)
        r.font.size = Pt(sz)
        r.italic = it
        if '———' in txt:
            r.font.color.rgb = RGBColor(0x55,0x55,0x55)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── 2. LOGOS (tableau 2 colonnes) ────────────────────────────────────────────
tbl_logo = doc.add_table(rows=1, cols=2)
tbl_logo.alignment = WD_TABLE_ALIGNMENT.CENTER

for col_idx, label in enumerate(["ISTAG", "IU DES TROPIQUES"]):
    cell = tbl_logo.rows[0].cells[col_idx]
    remove_cell_borders(cell)
    cell._tc.get_or_add_tcPr()
    # Hauteur de ligne 1.8 cm
    trPr = tbl_logo.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(Cm(1.8).pt * 20)))
    trPr.append(trHeight)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, 0, 0)
    r = p.add_run(f"[ Logo {label} ]")
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = BLEU

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── 3. BANDEAU PARCHEMIN (tableau 1 cellule fond bleu) ───────────────────────
tbl_titre = doc.add_table(rows=1, cols=1)
tbl_titre.alignment = WD_TABLE_ALIGNMENT.CENTER

cell_t = tbl_titre.rows[0].cells[0]
# Fond bleu
cell_shading(cell_t, '1F3A8A')
# Bordures de la cellule en bleu foncé
tc   = cell_t._tc
tcPr = tc.get_or_add_tcPr()
tcBorders = OxmlElement('w:tcBorders')
for side in ('top','left','bottom','right'):
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '12')
    b.set(qn('w:color'), '0A1F6E')
    tcBorders.append(b)
tcPr.append(tcBorders)

titre_lines = [
    "Développement d'une Application de",
    "Gestion de Voyages :",
    "Cas de FIRESOFTWARE",
]
for i, line in enumerate(titre_lines):
    p = cell_t.paragraphs[0] if i == 0 else cell_t.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, 60 if i == 0 else 0, 60 if i == len(titre_lines)-1 else 0)
    r = p.add_run(line)
    r.font.size = Pt(17)
    r.bold      = True
    r.italic    = True
    r.font.color.rgb = ORANGE

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── 4. DURÉE DU STAGE ────────────────────────────────────────────────────────
simple_para(doc, "Stage effectué du 12 Janvier au 12 Juin 2026",
            size=11, bold=True, before=0, after=40)

separator_para(doc, before=20, after=20)

# ── 5. DIPLÔME ───────────────────────────────────────────────────────────────
p_dip = doc.add_paragraph()
p_dip.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_spacing(p_dip, 20, 20)
add_run(p_dip, "En vue de l'obtention de la ", 11)
add_run(p_dip, "Licence Professionnelle (LP)", 11, bold=True, color=ORANGE)
add_run(p_dip, " en\n", 11)
add_run(p_dip, "Informatique de Gestion (IG)", 11, bold=True, color=ORANGE)

separator_para(doc, before=20, after=20)

# ── 6. ÉTUDIANT ──────────────────────────────────────────────────────────────
simple_para(doc, "Rédigé et présenté par", size=11, italic=True, before=20, after=20)
tiret_para(doc, before=10, after=10)
simple_para(doc, "YOUMBI FLORIDA", size=13, bold=True, before=0, after=30)
simple_para(doc, "Étudiante en Licence Professionnelle", size=10.5, before=0, after=60)

# ── 7. ENCADREURS (tableau 2 colonnes) ───────────────────────────────────────
tbl_enc = doc.add_table(rows=1, cols=2)
tbl_enc.alignment = WD_TABLE_ALIGNMENT.CENTER

enc_data = [
    [
        ("M. YAMSI CHRISTIAN",       10.5, True,  False),
        ("Encadreur Professionnel",   9.5,  False, True),
        ("Directeur Général de",      9.5,  False, False),
        ("FIRESOFTWARE",              9.5,  True,  False),
    ],
    [
        ("M. MBEKECK MARTIN ROLAND",  10.5, True,  False),
        ("Encadreur Académique",       9.5,  False, True),
        ("",                           9.5,  False, False),
        ("Enseignant à l'ISTAG",       9.5,  False, False),
    ],
]

for col_idx, lines in enumerate(enc_data):
    cell = tbl_enc.rows[0].cells[col_idx]
    remove_cell_borders(cell)
    for i, (txt, sz, bd, it) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, 0, 20)
        r = p.add_run(txt)
        r.font.size = Pt(sz)
        r.bold   = bd
        r.italic = it

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# ── 8. ANNÉE ACADÉMIQUE (tableau fond bleu) ───────────────────────────────────
tbl_an = doc.add_table(rows=1, cols=1)
tbl_an.alignment = WD_TABLE_ALIGNMENT.CENTER

# Largeur de la cellule à ~8 cm
tbl_an.columns[0].width = Cm(8)
cell_an = tbl_an.rows[0].cells[0]
cell_shading(cell_an, '1F3A8A')

# Bordures arrondies simulées (simple)
tc_an = cell_an._tc
tcPr_an = tc_an.get_or_add_tcPr()
tcBorders_an = OxmlElement('w:tcBorders')
for side in ('top','left','bottom','right'):
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '8')
    b.set(qn('w:color'), '0A1F6E')
    tcBorders_an.append(b)
tcPr_an.append(tcBorders_an)

for i, line in enumerate(["Année Académique", "2025 / 2026"]):
    p = cell_an.paragraphs[0] if i == 0 else cell_an.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p, 40 if i==0 else 0, 40 if i==1 else 0)
    r = p.add_run(line)
    r.font.size  = Pt(12)
    r.bold       = True
    r.font.color.rgb = BLANC

# ── Sauvegarde ────────────────────────────────────────────────────────────────
out = r"f:\Application\AgenceV\Page_Garde_v2.docx"
doc.save(out)
print(f"Fichier créé : {out}")
